"""
Convert all Markdown documentation files to Word (.docx) format.
Generates professionally formatted Word documents with proper headings,
tables, and styling for the Azani ISP Information System project.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOCS_DIR = os.path.join(os.path.dirname(__file__), 'docs')
OUTPUT_DIR = os.path.join(DOCS_DIR, 'word')

def create_styled_doc():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    
    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        h_style.font.bold = True
        if i == 1:
            h_style.font.size = Pt(22)
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(12)
        elif i == 2:
            h_style.font.size = Pt(16)
            h_style.paragraph_format.space_before = Pt(18)
            h_style.paragraph_format.space_after = Pt(8)
        elif i == 3:
            h_style.font.size = Pt(13)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
        else:
            h_style.font.size = Pt(11)
            h_style.paragraph_format.space_before = Pt(10)
            h_style.paragraph_format.space_after = Pt(4)
    
    return doc

def add_title_page(doc, title, subtitle="Azani Internet Service Provider Information System"):
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    run.font.name = 'Calibri'
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Prepared by: Derrick (Developer)")
    run3.font.size = Pt(12)
    run3.font.name = 'Calibri'
    
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Date: March 2026")
    run4.font.size = Pt(12)
    run4.font.name = 'Calibri'
    
    doc.add_page_break()

def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '003366',
        })
        borders.append(element)
    tbl_pr.append(borders)
    
    if table.rows:
        for cell in table.rows[0].cells:
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): '003366',
                qn('w:val'): 'clear',
            })
            shading.append(shading_elm)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(10)

def parse_markdown_to_docx(md_content, doc):
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                
                shading = p._element.get_or_add_pPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F0F0F0',
                    qn('w:val'): 'clear',
                })
                shading.append(shading_elm)
                
                code_lines = []
                in_code_block = False
            else:
                if in_table and table_rows:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            stripped = line.strip()
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells:
                in_table = True
                table_rows.append(cells)
            i += 1
            continue
        elif in_table and table_rows:
            _flush_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            text = _clean_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\s*\d+\.\s', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            text = _clean_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Number')
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            p = doc.add_paragraph()
            text = line.strip().strip('*').strip()
            run = p.add_run(text)
            run.bold = True
        elif line.strip() == '---':
            pass
        elif line.strip() == '':
            pass
        else:
            text = _clean_inline_formatting(line)
            if text.strip():
                doc.add_paragraph(text)
        
        i += 1
    
    if in_table and table_rows:
        _flush_table(doc, table_rows)

def _flush_table(doc, rows):
    if not rows:
        return
    
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append('')
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            text = _clean_inline_formatting(cell_text)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    
    style_table(table)
    doc.add_paragraph()

def _clean_inline_formatting(text):
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def convert_file(md_path, output_path, title):
    print(f"  Converting: {os.path.basename(md_path)} -> {os.path.basename(output_path)}")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = create_styled_doc()
    add_title_page(doc, title)
    parse_markdown_to_docx(content, doc)
    
    doc.save(output_path)

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    files_to_convert = [
        ('SRS/SRS.md', 'Software_Requirements_Specification.docx', 'Software Requirements Specification'),
        ('Design/SystemDesign.md', 'System_Design_Document.docx', 'System Design Document'),
        ('Architecture/Architecture.md', 'System_Architecture_Document.docx', 'System Architecture Document'),
        ('ProjectPlan.md', 'Project_Plan.docx', 'Project Plan Document'),
        ('Quotation.md', 'Quotation_Document.docx', 'Quotation Document'),
        ('TestingValidation.md', 'Testing_and_Validation.docx', 'Testing & Validation Document'),
        ('UserManual.md', 'User_Manual.docx', 'User Manual'),
        ('DeveloperNotes.md', 'Developer_Notes.docx', 'Developer Notes'),
        ('FrontendDesign.md', 'Frontend_Design_Document.docx', 'Frontend Design Document'),
        ('DatabaseDesign.md', 'Database_Design_Document.docx', 'Database Design Document'),
        ('DocumentationIndex.md', 'Documentation_Index.docx', 'Documentation Index'),
        ('DeploymentGuide.md', 'Deployment_Guide.docx', 'Deployment Guide'),
    ]
    
    print("=" * 60)
    print("  Azani ISP - Documentation Converter (MD -> DOCX)")
    print("=" * 60)
    print()
    
    converted = 0
    errors = []
    
    for md_file, docx_file, title in files_to_convert:
        md_path = os.path.join(DOCS_DIR, md_file)
        output_path = os.path.join(OUTPUT_DIR, docx_file)
        
        if os.path.exists(md_path):
            try:
                convert_file(md_path, output_path, title)
                converted += 1
            except Exception as e:
                errors.append((md_file, str(e)))
                print(f"  ERROR: {md_file} - {e}")
        else:
            print(f"  SKIP: {md_file} (not found)")
    
    print()
    print(f"  Converted: {converted}/{len(files_to_convert)}")
    if errors:
        print(f"  Errors: {len(errors)}")
        for f, e in errors:
            print(f"    - {f}: {e}")
    print(f"  Output: {OUTPUT_DIR}")
    print()

if __name__ == '__main__':
    main()
